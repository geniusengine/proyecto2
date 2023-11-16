import sys
from PyQt6.QtWidgets import QApplication, QMainWindow, QVBoxLayout, QWidget, QPushButton, QTextEdit, QFileDialog
from docx import Document
from docx2pdf import convert

class EstampadoApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.initUI()

    def initUI(self):
        self.setWindowTitle('Editor de Word')
        self.setGeometry(100, 100, 600, 400)

        self.central_widget = QWidget(self)
        self.setCentralWidget(self.central_widget)

        self.layout = QVBoxLayout()

        self.text_edit = QTextEdit(self)
        self.layout.addWidget(self.text_edit)

        self.btn_open = QPushButton('Abrir Documento Word', self)
        self.btn_open.clicked.connect(self.open_document)
        self.layout.addWidget(self.btn_open)

        self.btn_save = QPushButton('Guardar Documento Word', self)
        self.btn_save.clicked.connect(self.save_document)
        self.layout.addWidget(self.btn_save)

        self.btn_insert_image = QPushButton('Insertar Imagen', self)
        self.btn_insert_image.clicked.connect(self.insert_image)
        self.layout.addWidget(self.btn_insert_image)

        self.btn_back = QPushButton('Volver', self)
        self.btn_back.clicked.connect(self.close)
        self.layout.addWidget(self.btn_back)

        self.central_widget.setLayout(self.layout)

        self.current_document = Document()

    def open_document(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "Abrir Documento Word", "", "Documentos Word (*.docx)")

        if file_path:
            self.current_document = Document(file_path)
            content = ""
            for paragraph in self.current_document.paragraphs:
                content += paragraph.text + '\n'
            self.text_edit.setPlainText(content)

    def save_document(self):
        file_path, _ = QFileDialog.getSaveFileName(self, "Guardar Documento PDF", "", "Documentos PDF (*.pdf)")

        if file_path:
            # Guarda el documento de Word
            self.current_document.save(file_path + ".docx")

            # Convierte el documento de Word a PDF
            pdf_path = file_path + ".pdf"
            convert(file_path + ".docx", pdf_path)

    def insert_image(self):
        image_path, _ = QFileDialog.getOpenFileName(self, "Insertar Imagen", "", "Imágenes (*.png *.jpg *.bmp *.gif)")

        if image_path:
            # Agregar una imagen al documento Word
            paragraph = self.current_document.add_paragraph()
            run = paragraph.add_run()
            run.add_picture(image_path)

def main():
    app = QApplication(sys.argv)
    window = EstampadoApp()
    window.show()
    sys.exit(app.exec())

if __name__ == '__main__':
    main()
