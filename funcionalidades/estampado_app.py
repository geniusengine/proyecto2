import os
from PyQt6.QtWidgets import QApplication, QMainWindow, QPushButton, QVBoxLayout, QWidget
from PyQt6.QtGui import QIcon
from docx import Document
from docx2pdf import convert
from datetime import datetime
import tkinter as tk
from tkinter import filedialog
import logging
 

# Configurar el sistema de registro
logging.basicConfig(filename='registro.log', level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

class Estampadoxd(QMainWindow):
    def __init__(self, fechaNotificacion, numjui, nombTribunal, demandante, demandado, repre, mandante, domicilio, comuna, encargo, soli, arancel, parent=None):
        super().__init__(parent)

        self.setWindowTitle('Ventana de estampado')
        self.setWindowIcon(QIcon("static/icono-ventana.png"))
        self.setGeometry(100, 100, 400, 200)


        self.central_widget = QWidget(self)
        self.setCentralWidget(self.central_widget)

        # Layout vertical para organizar los botones
        layout_vertical = QVBoxLayout(self.central_widget)

        # Crear botones
        boton1 = QPushButton('Búsqueda Negativa 52', self)
        boton1.clicked.connect(self.estampar_1)

        boton2 = QPushButton('Búsqueda Negativa', self)
        boton2.clicked.connect(self.estampar_2)

        boton3 = QPushButton('Búsqueda Positivo', self)
        boton3.clicked.connect(self.estampar_3)

        boton4 = QPushButton('Búsqueda y Notificación', self)
        boton4.clicked.connect(self.estampar_4)

        boton5 = QPushButton('Notificación Personal', self)
        boton5.clicked.connect(self.estampar_5)

        boton6 = QPushButton('Embargo Vehículo', self)
        boton6.clicked.connect(self.estampar_6)

        boton7 = QPushButton('Acta Embargo', self)
        boton7.clicked.connect(self.estampar_7)

        boton8 = QPushButton('Embargo banco', self)
        boton8.clicked.connect(self.estampar_8)

        boton9 = QPushButton('Embargo muebles', self)
        boton9.clicked.connect(self.estampar_9)

        boton10 = QPushButton('Entrega material', self)
        boton10.clicked.connect(self.estampar_10)

        boton11 = QPushButton('Acta de Lanzamiento', self)
        boton11.clicked.connect(self.estampar_11)

        boton12 = QPushButton('Requerimiento de pago', self)
        boton12.clicked.connect(self.estampar_12)

        boton13 = QPushButton('Acta Embargo Frustrado', self)
        boton13.clicked.connect(self.estampar_13)

        boton14 = QPushButton('Inscripcion de embargo', self)
        boton14.clicked.connect(self.estampar_14)

        boton15 = QPushButton('Acta Embargo arrendamiento', self)
        boton15.clicked.connect(self.estampar_15)

        boton16 = QPushButton('Notificación subpersonal', self)
        boton16.clicked.connect(self.estampar_16)

        boton17 = QPushButton('Notificación por cedula', self)
        boton17.clicked.connect(self.estampar_17)

        boton18 = QPushButton('Notificación del conservador', self)
        boton18.clicked.connect(self.estampar_18)

        boton19 = QPushButton('Acta Opsicion al embargo', self)
        boton19.clicked.connect(self.estampar_19)

        boton20 = QPushButton('Acta pago por oposicion', self)
        boton20.clicked.connect(self.estampar_20)

        boton21= QPushButton('Retiro con fuerza pública', self)
        boton21.clicked.connect(self.estampar_21)


        # Agregar botones al layout vertical
        layout_vertical.addWidget(boton1)
        layout_vertical.addWidget(boton2)
        layout_vertical.addWidget(boton3)
        layout_vertical.addWidget(boton4)
        layout_vertical.addWidget(boton5)
        layout_vertical.addWidget(boton6)
        layout_vertical.addWidget(boton7)
        layout_vertical.addWidget(boton8)
        layout_vertical.addWidget(boton9)
        layout_vertical.addWidget(boton10)
        layout_vertical.addWidget(boton11)
        layout_vertical.addWidget(boton12)
        layout_vertical.addWidget(boton13)
        layout_vertical.addWidget(boton14)
        layout_vertical.addWidget(boton15)
        layout_vertical.addWidget(boton16)
        layout_vertical.addWidget(boton17)
        layout_vertical.addWidget(boton18)
        layout_vertical.addWidget(boton19)
        layout_vertical.addWidget(boton20)
        layout_vertical.addWidget(boton21)



        # Establecer el diseño principal de la ventana
        self.setLayout(layout_vertical)

        # Guardar los datos recibidos
        self.fechaNotificacion = fechaNotificacion
        self.numjui = numjui
        self.nombTribunal = nombTribunal
        self.demandante = demandante
        self.demandado = demandado
        self.repre = repre
        self.mandante = mandante
        self.domicilio = domicilio
        self.comuna = comuna
        self.encargo = encargo
        self.soli = soli
        self.arancel = arancel
        #comentarios ver si se pone
####################Hasta aqui hizo el bastian 1:30 am 05/01/2023-----------------------------------------------------------------------------------------------------------------------------

    #1 en este codigo cámbiame {self.nombdemandante} por {self.demandante} y los {self.apellidemandante} eliminalos
    def negativa52(self): 

        # variables de tiempo lel
        now = datetime.now()
        años = now.strftime("%d/%m/%y")
        horas = now.strftime("%H:%M")

        # Crea un nuevo documento de Word
        doc = Document()

        # Agrega el encabezado con los marcadores de posición
        encabezado = f"{self.nombTribunal}\n{self.numjui} : {self.encargo}\n {self.demandante} CON {self.demandado}"
        doc.add_paragraph(encabezado)

        # Agrega la sección de búsqueda negativa con los datos proporcionados
        busqueda_negativa = f"BÚSQUEDA NEGATIVA: Certifico haber buscado al(la) demandado(a) {self.demandado}, con domicilio en {self.domicilio} {self.comuna} especialmente el día {años}, siendo las {horas} horas, a fin de notificarle la resolución de fecha nose . Diligencia que no se llevó a efecto por cuanto el(la) demandado(a) no fue habido(a), {self.soli}. DOY FE."
        doc.add_paragraph(busqueda_negativa)

        # Agrega la firma al final del documento
        doc.add_paragraph(f"Drs. {self.arancel}.-")

        # Abre un cuadro de diálogo para seleccionar el directorio donde se guardará el documento
        root = tk.Tk()
        root.withdraw()  # Oculta la ventana de Tkinter
        save_path = filedialog.askdirectory()  # Abre el cuadro de diálogo

        # Guarda el documento en el directorio seleccionado
        doc.save(os.path.join(save_path, f'{self.numjui} {self.nombTribunal} {self.fechaNotificacion}.docx'))

        logging.info(f'A estamapado {self.numjui}-{self.encargo}')


        # Guarda el documento
        #desktop_path = os.path.expanduser('~')  # Obtiene el directorio del escritorio
        #doc.save(os.path.join(desktop_path, f'{self.numjui} {self.rolCausa}.docx'))

    #2
    def negativaP(self):

        # variables de tiempo lel
        now = datetime.now()
        años = now.strftime("%d/%m/%y")
        horas = now.strftime("%H:%M")

        # Crea un nuevo documento de Word
        doc = Document()

        # Agrega el encabezado con los marcadores de posición
        encabezado = f"{self.nombTribunal}\n{self.numjui}  :  {self.encargo}\n {self.demandante} CON {self.demandado}"
        doc.add_paragraph(encabezado)

        # Agrega la sección de búsqueda negativa con los datos proporcionados
        busqueda_negativaP = f"BÚSQUEDA NEGATIVA: Certifico haber buscado al(la) demandado(a) {self.demandado}, representado por {self.repre}, con domicilio en {self.domicilio} {self.comuna}, especialmente el día {años}, siendo las {horas} horas, a fin de notificarle la demanda íntegra y su respectivo proveído. Diligencia que no se llevó a efecto por cuanto el(la) demandado(a) no fue habido(a), {self.soli}. DOY FE."
        doc.add_paragraph(busqueda_negativaP)

        # Agrega la firma al final del documento
        doc.add_paragraph(f"Drs. {self.arancel}.-")

        # Abre un cuadro de diálogo para seleccionar el directorio donde se guardará el documento
        root = tk.Tk()
        root.withdraw()  # Oculta la ventana de Tkinter
        save_path = filedialog.askdirectory()  # Abre el cuadro de diálogo

        # Guarda el documento en el directorio seleccionado
        doc.save(os.path.join(save_path, f'{self.numjui} {self.nombTribunal} {self.fechaNotificacion}.docx'))

        logging.info(f'A estamapado {self.numjui}-{self.encargo}')

    #3
    def positivaP(self):# variables de tiempo lel
        now = datetime.now()
        años = now.strftime("%d/%m/%y")
        horas = now.strftime("%H:%M")
        
        # Crea un nuevo documento de Word
        doc = Document()

        # Agrega el encabezado con los marcadores de posición
        encabezado = f"{self.nombTribunal}\n{self.numjui}  :  {self.encargo}\n {self.demandante} CON {self.demandado}"
        doc.add_paragraph(encabezado)

        # Agrega la sección de búsqueda negativa con los datos proporcionados
        busqueda_positiva = f"BÚSQUEDA POSITIVA:a {años}, siendo las {horas} horas, en su domicilio ubicado en {self.domicilio} {self.comuna}, busqué a {self.demandado}, representado por {self.repre},a fin de notificarle la demanda íntegra y su respectivo proveído, diligencia que no se llevó a efecto por no ser habido en dicho domicilio, en ese momento. Por los dichos de una persona adulta, sexo masculino, vecino del lugar, se constató que este es el domicilio del demandado, y que se encuentra en el lugar del juicio. DOY FE. "
        doc.add_paragraph(busqueda_positiva)

        # Agrega la firma al final del documento
        doc.add_paragraph(f"Drs. {self.arancel}.-")

        # Abre un cuadro de diálogo para seleccionar el directorio donde se guardará el documento
        root = tk.Tk()
        root.withdraw()  # Oculta la ventana de Tkinter
        save_path = filedialog.askdirectory()  # Abre el cuadro de diálogo

         # Guarda el documento en el directorio seleccionado
        doc.save(os.path.join(save_path, f'{self.numjui} {self.nombTribunal} {self.fechaNotificacion}.docx'))

        logging.info(f'A estamapado {self.numjui}-{self.encargo}')
    
    

    #4
    def busquedaN(self):
                
        # variables de tiempo lel
        now = datetime.now()
        años = now.strftime("%d/%m/%y")
        horas = now.strftime("%H:%M")
        
        # Crea un nuevo documento de Word
        doc = Document()

        # Agrega el encabezado con los marcadores de posición
        encabezado = f"{self.nombTribunal}\n{self.numjui}  :  {self.encargo}\n {self.demandante} CON {self.demandado}"
        doc.add_paragraph(encabezado)

        # Agrega la sección de búsqueda negativa con los datos proporcionados
        busquedaN = f"BÚSQUEDA Y NOTIFICACIÓN: a {años}, siendo las {horas} horas, en su domicilio ubicado en {self.domicilio} {self.comuna}, busqué a {self.demandado} representado por {self.repre}, a fin de notificarle la resolución de fecha **, junto al escrito que antecede, diligencia que no se llevó a efecto por no ser habido en dicho domicilio, en ese momento. Informado por **, quien manifestó que es el domicilio del demandado y que se encuentra en el lugar del juicio, acto seguido procedo a notificar de conformidad al artículo 52 c.p.c. la resolución de fecha **, junto al escrito que antecede,  copia integra fue **. DOY FE."
        doc.add_paragraph(busquedaN)

        # Agrega la firma al final del documento
        doc.add_paragraph(f"Drs. {self.arancel}.-")

        # Abre un cuadro de diálogo para seleccionar el directorio donde se guardará el documento
        root = tk.Tk()
        root.withdraw()  # Oculta la ventana de Tkinter
        save_path = filedialog.askdirectory()  # Abre el cuadro de diálogo

         # Guarda el documento en el directorio seleccionado
        doc.save(os.path.join(save_path, f'{self.numjui} {self.nombTribunal} {self.fechaNotificacion}.docx'))

        logging.info(f'A estamapado {self.numjui}-{self.encargo}')


    #5
    def notificacionP(self):
                
        # variables de tiempo lel
        now = datetime.now()
        años = now.strftime("%d/%m/%y")
        horas = now.strftime("%H:%M")
        
        # Crea un nuevo documento de Word
        doc = Document()

        # Agrega el encabezado con los marcadores de posición
        encabezado = f"{self.nombTribunal}\n{self.numjui}  :  {self.encargo}\n {self.demandante} CON {self.demandado}"
        doc.add_paragraph(encabezado)

        # Agrega la sección de búsqueda negativa con los datos proporcionados
        notificacionP = f"NOTIFICACIÓN PERSONAL: En la Comuna de {self.comuna}, a {años}, siendo las {horas} horas, en su domicilio ubicado {self.domicilio}, {self.comuna}, notifiqué personalmente a {self.demandado}, Representado por {self.repre}, la demanda  y su respectivo proveído. Le hice entrega de copia íntegra de lo notificado y no firmó. La identidad del notificado/a  se estableció por los datos aportados por el mismo. Doy fe."
        doc.add_paragraph(notificacionP)

        # Agrega la firma al final del documento
        doc.add_paragraph(f"Drs. {self.arancel}.-") 

        # Abre un cuadro de diálogo para seleccionar el directorio donde se guardará el documento
        root = tk.Tk()
        root.withdraw()  # Oculta la ventana de Tkinter
        save_path = filedialog.askdirectory()  # Abre el cuadro de diálogo

         # Guarda el documento en el directorio seleccionado
        doc.save(os.path.join(save_path, f'{self.numjui} {self.nombTribunal} {self.fechaNotificacion}.docx'))

        logging.info(f'A estamapado {self.numjui}-{self.encargo}')

    #6
    def embargoV(self):
        now = datetime.now()
        fecha_actual = now.strftime("%d/%m/%y")
        hora_actual = now.strftime("%H:%M")

        # Crea un nuevo documento de Word
        doc = Document()

        # Agrega el encabezado con los marcadores de posición
        encabezado = f"JUZGADO= {self.nombTribunal}\nCAUSA ROL= {self.numjui} \nCARÁTULA {self.demandante} / {self.demandado}"
        doc.add_paragraph(encabezado)

        # Agrega la sección de acta de embargo
        acta_embargo = f"ACTA DE EMBARGO: En La Serena, a {fecha_actual} siendo las {hora_actual} horas, en mi oficio ubicado en calle Av. Del Mar, N° 5.700, of. N° 47 La Serena, a petición verbal del abogado doña NOMBRES APELLIDOS ABOGADA en representación de la parte ejecutante {self.demandante} RUT **, de conformidad a lo dispuesto en el artículo 447 del Código de Procedimiento Civil, procedo a trabar embargo sobre el(los) siguiente(s) bien(es) de propiedad del demandado {self.demandado}\n\n"
        acta_embargo += f"Tipo vehículo: \n"
        acta_embargo += f"Marca: \n"
        acta_embargo += f"Modelo: \n"
        acta_embargo += f"N° motor:\n"
        acta_embargo += f"N° Chassis: \n"
        acta_embargo += f"Color: \n"
        acta_embargo += f"PATENTE: \n"
        acta_embargo += f"Propietario: \n"
        acta_embargo += f"Rut: \n\n"
        acta_embargo += "Dicho vehículo quedó en poder del demandado en su calidad de depositario provisional, bajo su responsabilidad legal.\nCon lo actuado puse término a la diligencia levantando para constancia la presente acta. Doy fe."
        doc.add_paragraph(acta_embargo)

        # Agrega la firma al final del documento
        doc.add_paragraph(f"Drs. {self.arancel}.-")

        # Abre un cuadro de diálogo para seleccionar el directorio donde se guardará el documento
        root = tk.Tk()
        root.withdraw()  # Oculta la ventana de Tkinter
        save_path = filedialog.askdirectory()  # Abre el cuadro de diálogo

        # Guarda el documento en el directorio seleccionado
        doc.save(os.path.join(save_path, f'{self.numjui} {self.nombTribunal} {self.fechaNotificacion}.docx'))

        logging.info(f'A estamapado {self.numjui}-{self.encargo}')

    #7
    def actaEmbargo(self):
        # Variables de tiempo
        now = datetime.now()
        fecha_actual = now.strftime("%d/%m/%y")
        hora_actual = now.strftime("%H:%M")

        # Crea un nuevo documento de Word
        doc = Document()

        # Agrega el encabezado con los marcadores de posición
        encabezado = f"JUZGADO={self.nombTribunal}\nCAUSA ROL={self.numjui}  \n {self.demandado} / {self.demandante} "
        doc.add_paragraph(encabezado)

        # Agrega la sección de acta de embargo
        acta_embargo = f"ACTA DE EMBARGO: A FECHA {fecha_actual}, siendo las {hora_actual} horas, me constituí en el domicilio del demandado {self.demandado} representada legalmente por {self.repre} con domicilio ubicado en {self.domicilio}, {self.comuna}. Sin hacer uso de la fuerza pública, procedí a trabar embargo sobre:\n"
        acta_embargo += f"1. \n"
        acta_embargo += f"2. \n\n"
        acta_embargo += f"El referido bien embargado fue entregado al demandado, en forma simbólica, en su calidad de depositario provisional, bajo su responsabilidad legal.\n Drs. {self.arancel}.-"
        doc.add_paragraph(acta_embargo)

        # Abre un cuadro de diálogo para seleccionar el directorio donde se guardará el documento
        root = tk.Tk()
        root.withdraw()  # Oculta la ventana de Tkinter
        save_path = filedialog.askdirectory()  # Abre el cuadro de diálogo

        # Guarda el documento en el directorio seleccionado
        doc.save(os.path.join(save_path, f'{self.numjui} {self.nombTribunal} {self.fechaNotificacion}.docx'))

        logging.info(f'A estamapado {self.numjui}-{self.encargo}')

    
    def EmbargoBan(self):
        # Variables de tiempo
        now = datetime.now()
        fecha_actual = now.strftime("%d/%m/%y")
        hora_actual = now.strftime("%H:%M")

        # Crea un nuevo documento de Word
        doc = Document()

        # Agrega el encabezado con los marcadores de posición
        encabezado = f"{self.nombTribunal}\nROL/RIT: {self.numjui} \nCARATULADO: {self.encargo} CON {self.repre}"
        doc.add_paragraph(encabezado)

        # Agrega la sección de acta de embargo
        acta_embargo = f"ACTA DE EMBARGO: En la Comuna de {self.comuna}, a {fecha_actual}, siendo las {hora_actual} horas, en su domicilio ubicado en {self.domicilio}, notifiqué personalmente a {self.demandado} en representación del Banco **, solicitud y resolución de embargo por la suma de $** en la cuenta a nombre del demandado {self.demandado}. Le hice entrega de copia íntegra de lo notificado y firmó. La identidad del notificado se estableció por los datos aportados por el mismo. DOY FE.-\n\n"
        acta_embargo += "Señala que existían los dineros _________\nNo existen dineros en la cuenta _________  Cuenta cerrada _______\n"
        acta_embargo += f"Drs. Dist. Loc ${self.arancel}.-"
        doc.add_paragraph(acta_embargo)

        # Abre un cuadro de diálogo para seleccionar el directorio donde se guardará el documento
        root = tk.Tk()
        root.withdraw()  # Oculta la ventana de Tkinter
        save_path = filedialog.askdirectory()  # Abre el cuadro de diálogo

        # Guarda el documento en el directorio seleccionado
        doc.save(os.path.join(save_path, f'{self.numjui} {self.nombTribunal} {self.fechaNotificacion}.docx'))

        logging.info(f'A estamapado {self.numjui}-{self.encargo}')

    def EmbargoMue(self):
        # Variables de tiempo
        now = datetime.now()
        fecha_actual = now.strftime("%d/%m/%Y")

        # Crea un nuevo documento de Word
        doc = Document()

        # Agrega el encabezado con los marcadores de posición
        encabezado = f"{self.nombTribunal}\nCAUSA ROL= {self.encargo}\nCARATULA= {self.demandante}/ {self.demandado}"
        doc.add_paragraph(encabezado)

        # Agrega la sección de acta de embargo
        acta_embargo = f"ACTA DE EMBARGO: En La Serena, a {fecha_actual}, en mi oficio ubicado en Avenida de Mar No 5700, departamento No 47, Comuna de La Serena, a Solicitud del Abogado de la parte demandante procedo a trabar embargo sobre los bienes del demandado, consistente en el inmueble propiedad del deudor:\n\n"
        acta_embargo += f"Inmueble: **\n\n"
        acta_embargo += F"Drs. {self.arancel}.-"
        doc.add_paragraph(acta_embargo)

        # Abre un cuadro de diálogo para seleccionar el directorio donde se guardará el documento
        root = tk.Tk()
        root.withdraw()  # Oculta la ventana de Tkinter
        save_path = filedialog.askdirectory()  # Abre el cuadro de diálogo

        # Guarda el documento en el directorio seleccionado
        doc.save(os.path.join(save_path, f'{self.numjui} {self.nombTribunal} {self.fechaNotificacion}.docx'))

        logging.info(f'A estamapado {self.numjui}-{self.encargo}')

    def EntregaMaterial(self):
        # Variables de tiempo
        now = datetime.now()
        fecha_actual = now.strftime("%d/%m/%y")
        hora_actual = now.strftime("%H:%M")

        # Crea un nuevo documento de Word
        doc = Document()

        # Agrega el encabezado con los marcadores de posición
        encabezado = f"{self.nombTribunal}\nCAUSA ROL: {self.numjui} / {self.encargo}\nCARATULADO: {self.demandado} / {self.demandante}"
        doc.add_paragraph(encabezado)

        # Agrega la sección de acta de entrega material de inmueble
        acta_entrega_material = f"ACTA DE ENTREGA MATERIAL DE INMUEBLE: En {self.comuna}, a {fecha_actual}, siendo las {hora_actual} minutos. Me constituí en {self.domicilio}, {self.comuna}, según FECHA RESOLUCIÓN, con el fin de entregar la propiedad a la parte demandante la que se encontraba en las siguientes condiciones:\n\n"
        acta_entrega_material += "- \n- \n\n"
        acta_entrega_material += f"En general, la propiedad se encontraba en **. Se hizo entrega de la propiedad a la demandante siendo las 00:00 minutos, del día **. DOY FE.\n\n"
        acta_entrega_material += f"Drs. {self.arancel}."
        doc.add_paragraph(acta_entrega_material)

        # Abre un cuadro de diálogo para seleccionar el directorio donde se guardará el documento
        root = tk.Tk()
        root.withdraw()  # Oculta la ventana de Tkinter
        save_path = filedialog.askdirectory()  # Abre el cuadro de diálogo

        # Guarda el documento en el directorio seleccionado
        doc.save(os.path.join(save_path, f'{self.numjui} {self.nombTribunal} {self.fechaNotificacion}.docx'))

        logging.info(f'A estamapado {self.numjui}-{self.encargo}')
    
    def actaLanzamiento(self):
        # Variables de tiempohg{{´j}}
        now = datetime.now()
        fecha_actual = now.strftime("%d/%m/%Y")
        hora_actual = now.strftime("%H:%M")

        # Crea un nuevo documento de Word
        doc = Document()

        # Agrega el encabezado con los marcadores de posición
        encabezado = f"{self.nombTribunal}\nCAUSA ROL: {self.encargo}\nCARÁTULA: {self.demandante} / {self.demandado}"
        doc.add_paragraph(encabezado)

        # Agrega la sección de acta de lanzamiento
        acta_lanzamiento = f"ACTA DE LANZAMIENTO: En {self.comuna}, a {fecha_actual}, siendo las {hora_actual} hrs, me constituí en domicilio ubicado en {self.domicilio}, {self.comuna}, Región de Coquimbo, lugar donde procedí a efectuar el lanzamiento decretado en autos, terminando la actuación a las 00:00 hrs. Dejé constancia en fotografía, que se acompañan a continuación. Se hizo entrega de la propiedad a **. DOY FE.-\n\n"
        acta_lanzamiento += f"Drs. {self.arancel}.-"
        doc.add_paragraph(acta_lanzamiento)

        # Abre un cuadro de diálogo para seleccionar el directorio donde se guardará el documento
        root = tk.Tk()
        root.withdraw()  # Oculta la ventana de Tkinter
        save_path = filedialog.askdirectory()  # Abre el cuadro de diálogo

        # Guarda el documento en el directorio seleccionado
        doc.save(os.path.join(save_path, f'{self.numjui} {self.nombTribunal} {self.fechaNotificacion}.docx'))

        logging.info(f'A estamapado {self.numjui}-{self.encargo}')

    def actaRequerimientoPago(self):
        # Variables de tiempo
        now = datetime.now()
        fecha_actual = now.strftime("%d/%m/%Y")
        hora_actual = now.strftime("%H:%M")

        # Crea un nuevo documento de Word
        doc = Document()

        # Agrega el encabezado con los marcadores de posición
        encabezado = f"{self.nombTribunal}\n{self.numjui}: {self.encargo}\n {self.demandante} CON {self.demandado}"
        doc.add_paragraph(encabezado)

        # Agrega la sección de acta de requerimiento de pago en rebeldía
        acta_requerimiento_pago = f"ACTA DE REQUERIMIENTO DE PAGO EN REBELDÍA: En La Serena, a {fecha_actual}, siendo las {hora_actual} horas, en mi oficio ubicado en Av. Del Mar, N° 5.700, of. N° 47 La Serena, no habiendo comparecido el (la) demandado (a), don (ña) {self.demandado}, representado por {self.repre}, lo (la) doy por requerido(a) de pago en rebeldía por la suma de ** pesos (** pesos), más reajustes, intereses, recargos.\n\n"
        acta_requerimiento_pago += "NO SE EFECTUÓ EL PAGO.\n"
        acta_requerimiento_pago += "Tiene el (la) ejecutado (da) el plazo legal de 8 días hábiles para oponer excepciones a la ejecución.\n"
        acta_requerimiento_pago += "Para constancia levanté la presente acta. Doy fe.\n\n"
        acta_requerimiento_pago += f"Drs. + gastos. $ {self.arancel}.-"
        doc.add_paragraph(acta_requerimiento_pago)

        # Abre un cuadro de diálogo para seleccionar el directorio donde se guardará el documento
        root = tk.Tk()
        root.withdraw()  # Oculta la ventana de Tkinter
        save_path = filedialog.askdirectory()  # Abre el cuadro de diálogo

        # Guarda el documento en el directorio seleccionado
        doc.save(os.path.join(save_path, f'{self.numjui} {self.nombTribunal} {self.fechaNotificacion}.docx'))

        logging.info(f'A estamapado {self.numjui}-{self.encargo}')

    def actaEmbargoFrustrado(self):
        # Variables de tiempo
        now = datetime.now()
        fecha_actual = now.strftime("%d/%m/%Y")
        hora_actual = now.strftime("%H:%M")

        # Crea un nuevo documento de Word
        doc = Document()

        # Agrega el encabezado con los marcadores de posición
        encabezado = f"{self.nombTribunal} \n {self.numjui} : {self.encargo} \n {self.demandante} CON {self.demandado}"
        doc.add_paragraph(encabezado)

        # Agrega la sección de acta de embargo frustrado
        acta_embargo_frustrado = f"EMBARGO FRUSTRADO: En {self.comuna}, a {fecha_actual}, siendo las {hora_actual} horas, me apersoné en el domicilio de {self.demandado} ubicado en {self.domicilio}, {self.comuna}, con el fin de practicar el embargo ordenado en autos.\n"
        acta_embargo_frustrado += "No obstante, no puedo realizar la diligencia, persona adulta, de sexo femenino, permitió el ingreso, sin hacer uso de la fuerza pública, no obstante no pude llevar a efecto la diligencia por cuanto no hay bienes para la traba del embargo. DOY FE.-\n\n"
        acta_embargo_frustrado += f"Drs. {self.arancel}.-"
        doc.add_paragraph(acta_embargo_frustrado)

        # Abre un cuadro de diálogo para seleccionar el directorio donde se guardará el documento
        root = tk.Tk()
        root.withdraw()  # Oculta la ventana de Tkinter
        save_path = filedialog.askdirectory()  # Abre el cuadro de diálogo

        # Guarda el documento en el directorio seleccionado
        doc.save(os.path.join(save_path, f'{self.numjui} {self.nombTribunal} {self.fechaNotificacion}.docx'))

        logging.info(f'A estamapado {self.numjui}-{self.encargo}')

    def inscripcionEmbargo(self):
        # Variables de tiempo
        now = datetime.now()
        fecha_actual = now.strftime("%d/%m/%Y")
        hora_actual = now.strftime("%H:%M")

        # Crea un nuevo documento de Word
        doc = Document()

        # Agrega el encabezado con los marcadores de posición
        encabezado = f"{self.nombTribunal} \n{self.numjui} : {self.encargo }\n CON {self.demandado}"
        doc.add_paragraph(encabezado)

        # Agrega la sección de notificación de embargo
        notificacion_embargo = f"NOTIFICACIÓN: En la Comuna de La Serena, a {fecha_actual}, siendo las {hora_actual} horas, en su domicilio ubicado en Manuel Antonio Matta N°461, COMUNA DE LA SERENA, notifiqué personalmente a SERVICIO DE REGISTRO CIVIL E IDENTIFICACIÓN, solicitud de embargo en resolución de fecha **, del vehículo, marca, Modelo, número de motor, número de chasis, color, año, inscripción, inscrito a nombre de {self.demandado}, RUT **,. Le hice entrega de copia íntegra de lo notificado y firmó. La identidad del notificado se estableció por los datos aportados por el mismo, en el acto le requerí solicitud de embargo decretada en autos. DOY FE.-\n\n"
        notificacion_embargo += f"Drs. {self.arancel}.-"
        doc.add_paragraph(notificacion_embargo)

        # Abre un cuadro de diálogo para seleccionar el directorio donde se guardará el documento
        root = tk.Tk()
        root.withdraw()  # Oculta la ventana de Tkinter
        save_path = filedialog.askdirectory()  # Abre el cuadro de diálogo

        # Guarda el documento en el directorio seleccionado
        doc.save(os.path.join(save_path, f'{self.numjui} {self.nombTribunal} {self.fechaNotificacion}.docx'))

        logging.info(f'A estamapado {self.numjui}-{self.encargo}')

    def notificacionDesahucio(self):
        # Variables de tiempo
        now = datetime.now()
        fecha_actual = now.strftime("%d/%m/%Y")
        hora_actual = now.strftime("%H:%M")

        # Crea un nuevo documento de Word
        doc = Document()

        # Agrega el encabezado con los marcadores de posición
        encabezado = f"{self.nombTribunal}\n{self.numjui} : {self.encargo} \n  {self.demandante} CON {self.demandado}"
        doc.add_paragraph(encabezado)

        # Agrega la sección de notificación de desahucio
        notificacion_desahucio = f"En {self.comuna}, a {fecha_actual}, siendo las {hora_actual} HRS., en su domicilio ubicado en {self.domicilio}, {self.comuna}, Región de Coquimbo, notifiqué́ de conformidad al Artículo 8 No 2, y el artículo 18-A y artículo 18-C de la Ley 18.101, fijado por el Artículo 1o No 5 de la Ley 19.866 y en la forma que señala el inciso 2o del Artículo 44 del Código de Procedimiento Civil y lo dispuesto en la Ley 21.394, en relación con el Artículo 553 del mismo cuerpo, a {self.demandado}, la demanda y su respectivo proveído.\n\n"
        notificacion_desahucio += f"Acto seguido, lo requerí a fin de que dentro del plazo de diez días corridos, pague las rentas que se devenguen en el intertanto del plazo legal de desahucio, conforme lo pactado en el contrato de arrendamiento, no pudiendo imputar lo dado en garantía como precio ; Pagar todas las cuentas de consumo, gastos comunes y servicios adeudados hasta el día de la restitución del inmueble; Pagar los reajustes pactados de acuerdo al contrato de arrendamiento y los intereses de acuerdo a la ley; conjuntamente con los intereses y costas que correspondan, bajo apercibimiento de que si así no lo hiciere, no compareciera o no opusiera alguna excepción o defensa a la demanda dentro del plazo de 10 días corridos desde la fecha de notificación se le condenará de manera definitiva e invariable al pago de las rentas aludidas y asimismo se ordenará su LANZAMIENTO y de los demás ocupantes del inmueble arrendado, lo anterior dentro del plazo de veinte días desde la fecha de la notificación de la resolución de fecha 17 de marzo de 2023. La cédula conteniendo copia íntegra de lo notificado, se la entregué a persona adulta, sexo masculino que se encontraba en el lugar, quien no quiso dar su nombre.\n\n"
        notificacion_desahucio += f"Drs. {self.arancel}.-"
        doc.add_paragraph(notificacion_desahucio)

        # Abre un cuadro de diálogo para seleccionar el directorio donde se guardará el documento
        root = tk.Tk()
        root.withdraw()  # Oculta la ventana de Tkinter
        save_path = filedialog.askdirectory()  # Abre el cuadro de diálogo

        # Guarda el documento en el directorio seleccionado
        doc.save(os.path.join(save_path, f'{self.numjui} {self.nombTribunal} {self.fechaNotificacion}.docx'))

        logging.info(f'A estamapado {self.numjui}-{self.encargo}')
    
    def notificacionPersonalSubsidiaria(self):
        # Variables de tiempo
        now = datetime.now()
        fecha_actual = now.strftime("%d/%m/%Y")
        hora_actual = now.strftime("%H:%M")

        # Crea un nuevo documento de Word
        doc = Document()

        # Agrega el encabezado con los marcadores de posición
        encabezado = f"{self.nombTribunal}\n{self.numjui} : {self.encargo} \n {self.demandante} CON {self.demandado}"
        doc.add_paragraph(encabezado)

        # Agrega la sección de notificación personal subsidiaria
        notificacion_subsidiaria = f"NOTIFICACIÓN PERSONAL SUBSIDIARIA: A {fecha_actual}, siendo las {hora_actual} horas, en su domicilio ubicado en {self.domicilio}, {self.comuna}, notifiqué a {self.demandado}, representado por {self.repre}, se constató que es su domicilio y que se encuentra en el lugar del juicio por **, en conformidad al artículo 44 inciso segundo del Código de Procedimiento Civil modificado por el artículo 3, N°3 letra A y B de la ley N°21.394, la demanda íntegra, su respectivo proveído. Copia íntegra de lo notificado fue **. Doy fe.\n"
        notificacion_subsidiaria += f"Búsqueda: **.-\nDrs: {self.arancel}.-"
        doc.add_paragraph(notificacion_subsidiaria)

        # Agrega la sección de cédula de espera
        cedula_espera = f"CÉDULA DE ESPERA: A **, siendo las ** horas, en su domicilio ubicado en {self.domicilio}, {self.comuna}, busqué a {self.demandado}, representado por {self.repre}, y notifiqué el mandamiento de fs.01, por cédula. Dejé citado (a) al demandado (a) a fin de que concurra a mi oficio ubicado en calle Av. Del Mar, N° 5.700, of. N° 47 La Serena., el día **, a las ** horas, a fin de requerirlo (a) personalmente de pago, bajo apercibimiento de que si no comparece, será declarado (a) requerido (a) de pago en rebeldía y se le practicará sin más trámite el embargo.\n"
        cedula_espera += "Para constancia levanté la presente acta. Doy fe.-"
        doc.add_paragraph(cedula_espera)

        # Abre un cuadro de diálogo para seleccionar el directorio donde se guardará el documento
        root = tk.Tk()
        root.withdraw()  # Oculta la ventana de Tkinter
        save_path = filedialog.askdirectory()  # Abre el cuadro de diálogo

        # Guarda el documento en el directorio seleccionado
        doc.save(os.path.join(save_path, f'{self.numjui} {self.nombTribunal} {self.fechaNotificacion}.docx'))

        logging.info(f'A estamapado {self.numjui}-{self.encargo}')

    def notificacionPorCedula(self):
            # Variables de tiempo
            now = datetime.now()
            fecha_actual = now.strftime("%d/%m/%Y")
            hora_actual = now.strftime("%H:%M")

            # Crea un nuevo documento de Word
            doc = Document()

            # Agrega el encabezado con los marcadores de posición
            encabezado = f"{self.nombTribunal}\n{self.numjui} : {self.encargo} \n {self.demandante} CON {self.demandado}"
            doc.add_paragraph(encabezado)

            # Agrega la sección de notificación por cédula
            notificacion_cedula = f"NOTIFICACIÓN POR CÉDULA: a {fecha_actual}, siendo las {hora_actual} horas, en su domicilio ubicado en {self.domicilio}, {self.comuna}, notifiqué por cédula a {self.demandado} representado por {self.repre}, a fin de notificarle la resolución de fecha **, la resolución de fecha **. La cédula conteniendo copia íntegra de lo notificado fue **. DOY FE.\n"
            notificacion_cedula += f"Drs. {self.arancel}.-"
            doc.add_paragraph(notificacion_cedula)

            # Abre un cuadro de diálogo para seleccionar el directorio donde se guardará el documento
            root = tk.Tk()
            root.withdraw()  # Oculta la ventana de Tkinter
            save_path = filedialog.askdirectory()  # Abre el cuadro de diálogo

            # Guarda el documento en el directorio seleccionado
            doc.save(os.path.join(save_path, f'{self.numjui} {self.nombTribunal} {self.fechaNotificacion}.docx'))

            logging.info(f'A estamapado {self.numjui}-{self.encargo}')

    def notificacionConsevador(self):
        # Variables de tiempo
        now = datetime.now()
        fecha_actual = now.strftime("%d/%m/%Y")
        hora_actual = now.strftime("%H:%M")

        # Crea un nuevo documento de Word
        doc = Document()

        # Agrega el encabezado con los marcadores de posición
        encabezado = f"{self.nombTribunal}\n{self.numjui} : {self.encargo} \n {self.demandante} CON {self.demandado}"
        doc.add_paragraph(encabezado)

        # Agrega la sección de notificación
        notificacion_texto = f"NOTIFICACIÓN: En la comuna de {self.comuna}, a {fecha_actual}, siendo las {hora_actual} horas, en su domicilio ubicado en {self.domicilio}, notifiqué personalmente a don {self.demandado}, solicitud de **. Le hice entrega de copia íntegra de lo notificado y firmó. La identidad del notificado se estableció por los datos aportados por el mismo, en el acto le requerí solicitud de **. DOY FE.-\n"
        notificacion_texto += f"Drs. {self.arancel}.-"
        doc.add_paragraph(notificacion_texto)

        # Abre un cuadro de diálogo para seleccionar el directorio donde se guardará el documento
        root = tk.Tk()
        root.withdraw()  # Oculta la ventana de Tkinter
        save_path = filedialog.askdirectory()  # Abre el cuadro de diálogo

        # Guarda el documento en el directorio seleccionado
        doc.save(os.path.join(save_path, f'{self.numjui} {self.nombTribunal} {self.fechaNotificacion}.docx'))

        logging.info(f'A estamapado {self.numjui}-{self.encargo}')
        
    def actaOposicionEmbargo(self):
        # Variables de tiempo
        now = datetime.now()
        fecha_actual = now.strftime("%d/%m/%Y")
        hora_actual = now.strftime("%H:%M")

        # Crea un nuevo documento de Word
        doc = Document()

        # Agrega el encabezado con los marcadores de posición
        encabezado = f"{self.nombTribunal}\n{self.numjui} : {self.encargo} \n {self.demandante} CON {self.demandado}"
        doc.add_paragraph(encabezado)

        # Agrega la sección de acta de oposición a embargo
        acta_oposicion_embargo = f"ACTA DE OPOSICION A EMBARGO: a {fecha_actual}, siendo las {hora_actual} horas, concurrí al domicilio ubicado en {self.domicilio}, {self.comuna}, notifiqué a {self.demandado}, representado por {self.repre}, con el objeto de proceder a trabar embargo sobre bienes de su propiedad, diligencia que no pude efectuar, por OPOSICIÓN AL EMBARGO, por **. DOY FE.\n"
        acta_oposicion_embargo += f"Drs. {self.arancel}.-"
        doc.add_paragraph(acta_oposicion_embargo)

        # Abre un cuadro de diálogo para seleccionar el directorio donde se guardará el documento
        root = tk.Tk()
        root.withdraw()  # Oculta la ventana de Tkinter
        save_path = filedialog.askdirectory()  # Abre el cuadro de diálogo

        # Guarda el documento en el directorio seleccionado
        doc.save(os.path.join(save_path, f'{self.numjui} {self.nombTribunal} {self.fechaNotificacion}.docx'))

        logging.info(f'A estamapado {self.numjui}-{self.encargo}')

    def requerimientoPagoOposicionEmbargo(self):
        # Variables de tiempo
        now = datetime.now()
        fecha_actual = now.strftime("%d/%m/%Y")
        hora_actual = now.strftime("%H:%M")

        # Crea un nuevo documento de Word
        doc = Document()

        # Agrega el encabezado con los marcadores de posición
        encabezado = f"{self.nombTribunal}\n{self.numjui} : {self.encargo} \n {self.demandante} CON {self.demandado}"
        doc.add_paragraph(encabezado)

        # Agrega la sección de requerimiento de pago y oposición a embargo
        requerimiento_oposicion_texto = f"REQUERIMIENTO DE PAGO PERSONAL Y OPOSICIÓN A EMBARGO: A {fecha_actual}, siendo las {hora_actual} horas, en su domicilio ubicado en {self.domicilio}, {self.comuna}, notifiqué el mandamiento de **, y requerí personalmente de pago a {self.demandado}, representado por {self.repre}, a fin de que en el acto pague a {self.demandante}, o a quien sus derechos represente, la suma de ** (** pesos), más reajustes, intereses y recargos.\n"
        requerimiento_oposicion_texto += f"NO SE EFECTUÓ EL PAGO. Le hice presente la orden de embargo sobre bienes de su propiedad, en cantidad suficiente que permita cubrir el valor total de lo adeudado, más los gastos de la ejecución. También le notifiqué su designación como depositario provisional de los bienes que se le embarguen, cargo que aceptó, jurando su fiel desempeño y además le hice presente el plazo legal que tiene para deducir excepciones a la ejecución.\n"
        requerimiento_oposicion_texto += f"Acto seguido intenté trabar embargo sobre bienes de su propiedad, diligencia que no se llevó a efecto por OPOSICIÓN DEL PROPIO DEMANDADO.\n"
        requerimiento_oposicion_texto += f"Para constancia levanté la presente acta que la ejecutada no firmó, entregándole en el acto copia del mandamiento. Doy Fe."
        doc.add_paragraph(requerimiento_oposicion_texto)

        # Abre un cuadro de diálogo para seleccionar el directorio donde se guardará el documento
        root = tk.Tk()
        root.withdraw()  # Oculta la ventana de Tkinter
        save_path = filedialog.askdirectory()  # Abre el cuadro de diálogo

        # Guarda el documento en el directorio seleccionado
        doc.save(os.path.join(save_path, f'{self.numjui} {self.nombTribunal} {self.fechaNotificacion}.docx'))

        logging.info(f'A estamapado {self.numjui}-{self.encargo}')

    def retiroVehiculo(self):
        # Variables de tiempo
        now = datetime.now()
        fecha_actual = now.strftime("%d/%m/%Y")
        hora_actual = now.strftime("%H:%M")

        # Crea un nuevo documento de Word
        doc = Document()

        # Agrega el encabezado con los marcadores de posición
        encabezado = f"{self.nombTribunal}\n{self.numjui} : {self.encargo} \n {self.demandante} CON {self.demandado}"
        doc.add_paragraph(encabezado)

        # Agrega la sección de retiro de vehículo
        retiro_vehiculo_texto = f"En La Serena, a {fecha_actual}, siendo las {hora_actual} horas, me constituí en {self.domicilio}, comuna de {self.comuna}, en presencia de **, placa N° **, cumpliendo la resolución de fecha **, para ser RETIRO DE VEHÍCULO CON FUERZA PÚBLICA Y GRUA, correspondiente a:\n"
        retiro_vehiculo_texto += "Tipo vehículo: \nMarca: \nModelo: \nN° motor: \nN° Chassis: \nColor: \nPATENTE: \nPropietario: \nRut: }\n"
        retiro_vehiculo_texto += f"Se retiró ante persona adulta, sexo masculino, de nombre **, Rut.** , en **.\n"
        retiro_vehiculo_texto += f"Se adjuntan fotografías para mayor detalle del vehículo, encontrado en buen estado, se entregó copia del oficio a **, Rut. **. Se hizo entrega del vehículo a la parte demandante en su domicilio ubicado en **-**. DOY FE.-\n"
        retiro_vehiculo_texto += f"DRS. {self.arancel}."
        doc.add_paragraph(retiro_vehiculo_texto)

        # Abre un cuadro de diálogo para seleccionar el directorio donde se guardará el documento
        root = tk.Tk()
        root.withdraw()  # Oculta la ventana de Tkinter
        save_path = filedialog.askdirectory()  # Abre el cuadro de diálogo

        # Guarda el documento en el directorio seleccionado
        doc.save(os.path.join(save_path, f'{self.numjui} {self.nombTribunal} {self.fechaNotificacion}.docx'))

        logging.info(f'A estamapado {self.numjui}-{self.encargo}')

    # Botones enlazados con sus opciones correspindientes    
    def estampar_1(self):
        self.negativa52()

    def estampar_2(self):
        self.negativaP()

    def estampar_3(self):
        self.positivaP()

    def estampar_4(self):
        self.busquedaN()

    def estampar_5(self):
        self.notificacionP()

    def estampar_6(self):
        self.embargoV()
    
    def estampar_7(self):
        self.actaEmbargo()

    def estampar_8(self):
        self.EmbargoBan()

    def estampar_9(self):
        self.EmbargoMue()

    def estampar_10(self):
        self.EntregaMaterial()

    def estampar_11(self):
        self.actaLanzamiento()
        
    def estampar_12(self):
        self.actaRequerimientoPago()
    
    def estampar_13(self):
        self.actaEmbargoFrustrado()

    def estampar_14(self):
        self.inscripcionEmbargo()
        
    def estampar_15(self):
        self.notificacionDesahucio()

    def estampar_16(self):
        self.notificacionPersonalSubsidiaria()

    def estampar_17(self):
        self.notificacionPorCedula()
        
    def estampar_18(self):
        self.notificacionConsevador()

    def estampar_19(self):
        self.actaOposicionEmbargo()

    def estampar_20(self):
        self.requerimientoPagoOposicionEmbargo()
        
    def estampar_21(self):
        self.retiroVehiculo()



def main():
    app = QApplication([])
    ventana = Estampadoxd()
    ventana.show()
    app.exec()

if __name__ == '__main__':
    main()
