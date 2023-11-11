import sys
from PyQt6.QtWidgets import QApplication, QMainWindow, QVBoxLayout, QWidget, QPushButton, QTextEdit, QFileDialog
from PyQt6.QtGui import QTextDocument, QImage, QPainter
from PyQt6.QtPrintSupport import QPrinter
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

class EstampadoApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.initUI()

    def initUI(self):
        self.setWindowTitle('Editor de Word')
        self.setGeometry(100, 100, 600, 400)

        self.central_widget = QWidget(self)
        self.setCentralWidget(self.central_widget)

        self.layout = QVBoxLayout()

        self.text_edit = QTextEdit(self)
        self.layout.addWidget(self.text_edit)

        self.btn_open = QPushButton('Abrir Documento', self)
        self.btn_open.clicked.connect(self.open_document)
        self.layout.addWidget(self.btn_open)

        self.btn_save_pdf = QPushButton('Guardar como PDF', self)
        self.btn_save_pdf.clicked.connect(self.save_as_pdf)
        self.layout.addWidget(self.btn_save_pdf)

        # Botón para insertar imagen
        self.btn_insert_image = QPushButton('Insertar Imagen', self)
        self.btn_insert_image.clicked.connect(self.insert_image)
        self.layout.addWidget(self.btn_insert_image)

        self.btn_back = QPushButton('Volver', self)
        self.btn_back.clicked.connect(self.close)
        self.layout.addWidget(self.btn_back)

        self.central_widget.setLayout(self.layout)

        self.current_document_path = None

    def open_document(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "Abrir Documento Word", "", "Documentos Word (*.docx)")

        if file_path:
            self.current_document_path = file_path
            document = Document(file_path)
            content = ""
            for paragraph in document.paragraphs:
                content += paragraph.text + '\n'
            self.text_edit.setPlainText(content)

    def save_as_pdf(self):
        if self.current_document_path:
            pdf_file_path, _ = QFileDialog.getSaveFileName(self, "Guardar como PDF", "", "Archivos PDF (*.pdf)")

            if pdf_file_path:
                # Crear un objeto QTextDocument y establecer su contenido
                text_document = QTextDocument()
                text_document.setPlainText(self.text_edit.toPlainText())

                # Crear una imagen del QTextDocument
                image = QImage(text_document.size().toSize(), QImage.Format_ARGB32)
                painter = QPainter(image)
                text_document.drawContents(painter)
                painter.end()

                # Crear un objeto Canvas de reportlab para el PDF
                pdf_canvas = canvas.Canvas(pdf_file_path, pagesize=letter)
                width, height = letter
                pdf_canvas.drawInlineImage(image, 0, 0, width, height)

                # Cerrar el objeto Canvas
                pdf_canvas.save()

                self.statusBar().showMessage("Documento guardado como PDF con éxito.")
        else:
            self.statusBar().showMessage("Abre o guarda un documento Word antes de convertirlo a PDF.")

    def insert_image(self):
        # Abre un cuadro de diálogo para seleccionar una imagen
        image_path, _ = QFileDialog.getOpenFileName(self, "Insertar Imagen", "", "Imágenes (*.png *.jpg *.bmp *.gif)")

        if image_path:
            # Inserta la imagen en el QTextEdit
            cursor = self.text_edit.textCursor()
            cursor.insertImage(image_path)

def main():
    app = QApplication(sys.argv)
    window = EstampadoApp()
    window.show()
    sys.exit(app.exec())

if __name__ == '__main__':
    main()
